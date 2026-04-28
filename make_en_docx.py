import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Color Palette (LA Market Professional) ───
COLOR_DARK    = RGBColor(0x1A, 0x1A, 0x2E)   # deep navy
COLOR_ACCENT  = RGBColor(0x00, 0x78, 0xD0)   # vivid blue
COLOR_RED     = RGBColor(0xC0, 0x00, 0x00)   # important notice red
COLOR_GRAY    = RGBColor(0x60, 0x60, 0x60)   # body text gray
COLOR_LIGHT   = RGBColor(0xF5, 0xF8, 0xFF)  # light blue bg
COLOR_BORDER  = RGBColor(0xCC, 0xCC, 0xCC)   # table border

FONT_BODY    = "Arial"
FONT_HEADING = "Arial"
FONT_MONO    = "Courier New"

# ─── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), 'single')
        tag.set(qn('w:sz'), '4')
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), kwargs.get('color','CCCCCC'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def set_para_border_bottom(para, color="0078D0", sz="8"):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), sz)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_horizontal_rule(doc, color="0078D0"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    p._p.append(pBdr)
    return p

def para_style(para, size, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
               space_before=0, space_after=6, italic=False):
    para.alignment = align
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(space_after)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    para.paragraph_format.line_spacing = 1.15
    for run in para.runs:
        run.font.size  = Pt(size)
        run.font.bold  = bold
        run.font.italic = italic
        run.font.name  = FONT_BODY
        if color:
            run.font.color.rgb = color

# ─── Core parser ───────────────────────────────────────────────────────────────

def md_to_docx(md_path, docx_path, title=""):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.read().split('\n')

    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(3.0)

    # Default style
    style = doc.styles['Normal']
    style.font.name = FONT_BODY
    style.font.size = Pt(10.5)

    i = 0
    while i < len(lines):
        line = lines[i]

        # ═══════════════════════════════════════════════════
        # H1  ─── Title (centered, large, accent underline)
        # ═══════════════════════════════════════════════════
        if line.startswith('# ') and not line.startswith('# AI'):
            text = line[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after  = Pt(4)
            run = p.add_run(text)
            run.font.size  = Pt(22)
            run.font.bold  = True
            run.font.color.rgb = COLOR_DARK
            run.font.name  = FONT_HEADING
            i += 1
            continue

        # H2 ─── Article / Section heading
        if line.startswith('## '):
            text = line[3:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(text)
            run.font.size  = Pt(13)
            run.font.bold  = True
            run.font.color.rgb = COLOR_ACCENT
            run.font.name  = FONT_HEADING
            set_para_border_bottom(p, color="0078D0", sz="4")
            i += 1
            continue

        # H3 ─── Sub-section
        if line.startswith('### '):
            text = line[4:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(text)
            run.font.size  = Pt(11)
            run.font.bold  = True
            run.font.color.rgb = COLOR_DARK
            run.font.name  = FONT_HEADING
            i += 1
            continue

        # H4 ─── Minor heading
        if line.startswith('#### '):
            text = line[5:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(text)
            run.font.size  = Pt(10.5)
            run.font.bold  = True
            run.font.color.rgb = COLOR_GRAY
            run.font.name  = FONT_HEADING
            i += 1
            continue

        # ═══════════════════════════════════════════════════
        # TABLES
        # ═══════════════════════════════════════════════════
        if line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1
            # Parse
            rows_data = []
            for j, tl in enumerate(table_lines):
                if j == 1: continue  # skip separator |---|---|
                cells = [c.strip() for c in tl.strip('|').split('|')]
                rows_data.append(cells)
            if not rows_data: continue

            cols = len(rows_data[0])
            tbl  = doc.add_table(rows=len(rows_data), cols=cols)
            tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
            tbl.style = 'Table Grid'

            for ri, row_data in enumerate(rows_data):
                is_header = ri == 0
                for ci, cell_text in enumerate(row_data):
                    cell = tbl.rows[ri].cells[ci]
                    cell.text = ''
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(3)
                    p.paragraph_format.space_after  = Pt(3)
                    run = p.add_run(cell_text)
                    run.font.name  = FONT_BODY
                    run.font.size  = Pt(10)
                    if is_header:
                        run.font.bold  = True
                        run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
                        set_cell_bg(cell, '1A1A2E')
                    else:
                        run.font.color.rgb = COLOR_DARK
                        bg = 'F5F8FF' if ri % 2 == 0 else 'FFFFFF'
                        set_cell_bg(cell, bg)
                    set_cell_border(cell, color='CCCCCC')
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
            continue

        # ═══════════════════════════════════════════════════
        # BLOCKQUOTE / IMPORTANT NOTICE
        # ═══════════════════════════════════════════════════
        if line.startswith('>'):
            text = line.lstrip('>').strip()
            container = doc.add_paragraph()
            container.paragraph_format.left_indent  = Inches(0.3)
            container.paragraph_format.right_indent = Inches(0.3)
            container.paragraph_format.space_before = Pt(8)
            container.paragraph_format.space_after  = Pt(8)
            run = container.add_run(text)
            run.font.size   = Pt(10.5)
            run.font.bold   = True
            run.font.italic = True
            run.font.color.rgb = COLOR_RED
            run.font.name   = FONT_BODY
            # Add left border accent
            pPr = container._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '16')
            left.set(qn('w:space'), '4')
            left.set(qn('w:color'), 'C00000')
            pBdr.append(left)
            pPr.append(pBdr)
            i += 1
            continue

        # ═══════════════════════════════════════════════════
        # HORIZONTAL RULE
        # ═══════════════════════════════════════════════════
        if line.strip() in ('---', '***', '___'):
            add_horizontal_rule(doc)
            i += 1
            continue

        # ═══════════════════════════════════════════════════
        # BULLET LIST
        # ═══════════════════════════════════════════════════
        if line.startswith('- [ ]') or line.startswith('- '):
            if line.startswith('- [ ]'):
                text = line[5:].strip()
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.left_indent  = Inches(0.25)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after  = Pt(2)
                run = p.add_run('☐  ' + text)
                run.font.name  = FONT_BODY
                run.font.size  = Pt(10.5)
                run.font.color.rgb = COLOR_DARK
            else:
                text = line[2:].strip()
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.left_indent  = Inches(0.25)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after  = Pt(2)
                # Bold (**text**) support
                parts = re.split(r'\*\*(.*?)\*\*', text)
                for pi, part in enumerate(parts):
                    run = p.add_run(part)
                    run.font.name  = FONT_BODY
                    run.font.size  = Pt(10.5)
                    run.font.color.rgb = COLOR_DARK
                    if pi % 2 == 1:
                        run.font.bold = True
            i += 1
            continue

        # ═══════════════════════════════════════════════════
        # NUMBERED LIST
        # ═══════════════════════════════════════════════════
        if re.match(r'^\d+\.\s', line):
            m = re.match(r'^(\d+)\.\s+(.*)', line)
            num, text = m.group(1), m.group(2)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.left_indent  = Inches(0.3)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            parts = re.split(r'\*\*(.*?)\*\*', text)
            for pi, part in enumerate(parts):
                run = p.add_run(part)
                run.font.name  = FONT_BODY
                run.font.size  = Pt(10.5)
                run.font.color.rgb = COLOR_DARK
                if pi % 2 == 1:
                    run.font.bold = True
            i += 1
            continue

        # ═══════════════════════════════════════════════════
        # EMPTY LINE
        # ═══════════════════════════════════════════════════
        if not line.strip():
            i += 1
            continue

        # ═══════════════════════════════════════════════════
        # PARAGRAPH — with **bold** support
        # ═══════════════════════════════════════════════════
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(4)
        parts = re.split(r'\*\*(.*?)\*\*', line)
        for pi, part in enumerate(parts):
            run = p.add_run(part)
            run.font.name  = FONT_BODY
            run.font.size  = Pt(10.5)
            run.font.color.rgb = COLOR_DARK
            if pi % 2 == 1:
                run.font.bold = True
        i += 1

    doc.save(docx_path)
    size = os.path.getsize(docx_path) // 1024
    print(f"[OK] {os.path.basename(docx_path):50s}  ({size} KB)")

# ─── Run ────────────────────────────────────────────────────────────────────────

base = r'I:\Portraitpay ai\Doc\合同模板设计\EN'
files = sorted([
    ('00-Overview-Guide.docx.md',                    '00-Overview-and-Signing-Guide.docx'),
    ('01-Standard-License-Agreement.docx.md',         '01-Standard-License-Agreement.docx'),
    ('02-Exclusive-License-Agreement.docx.md',        '02-Exclusive-License-Agreement.docx'),
    ('03-Endorsement-License-Agreement.docx.md',      '03-Endorsement-License-Agreement.docx'),
    ('04-Film-Adaptation-License-Agreement.docx.md',  '04-Film-Adaptation-License-Agreement.docx'),
])

for md_name, docx_name in files:
    md_path  = os.path.join(base, md_name)
    docx_path = os.path.join(base, docx_name)
    md_to_docx(md_path, docx_path)

print("\nAll English Word documents generated successfully!")
