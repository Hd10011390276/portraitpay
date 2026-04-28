import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

def md_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # Set default font for Chinese
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        # Heading 1
        if line.startswith('# '):
            h = doc.add_heading(line[2:], level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Heading 2
        elif line.startswith('## '):
            h = doc.add_heading(line[3:], level=2)
        # Heading 3
        elif line.startswith('### '):
            h = doc.add_heading(line[4:], level=3)
        # Table
        elif line.startswith('|'):
            # Collect table rows
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1
            # Parse table (skip separator row index 1)
            rows_data = []
            for j, tl in enumerate(table_lines):
                if j == 1:  # separator row
                    continue
                cells = [c.strip() for c in tl.strip('|').split('|')]
                rows_data.append(cells)
            if rows_data:
                cols = len(rows_data[0])
                table = doc.add_table(rows=len(rows_data), cols=cols)
                table.style = 'Table Grid'
                for ri, row_data in enumerate(rows_data):
                    for ci, cell_text in enumerate(row_data):
                        cell = table.rows[ri].cells[ci]
                        cell.text = cell_text
                doc.add_paragraph()
                continue
        # Blockquote / important note
        elif line.startswith('>'):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(line[1:].strip())
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
            run.font.bold = True
        # Bullet list
        elif line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            # Handle bold in bullet
            parts = re.split(r'\*\*(.*?)\*\*', line[2:])
            for pi, part in enumerate(parts):
                if pi % 2 == 1:
                    run = p.add_run(part)
                    run.bold = True
                else:
                    p.add_run(part)
        # Separator
        elif line.startswith('---'):
            p = doc.add_paragraph('─' * 50)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Empty line
        elif line.strip() == '':
            pass
        # Regular paragraph
        else:
            p = doc.add_paragraph()
            # Handle bold/italic
            parts = re.split(r'\*\*(.*?)\*\*', line)
            for pi, part in enumerate(parts):
                if pi % 2 == 1:
                    run = p.add_run(part)
                    run.bold = True
                else:
                    p.add_run(part)

        i += 1

    doc.save(docx_path)
    print(f"Saved: {docx_path}")

# Convert audit report
md_report = r'C:\Users\Administrator\.openclaw\workspace\PortraitPay\audit\2026-04-28-frontend-audit.md'
docx_report = r'C:\Users\Administrator\.openclaw\workspace\PortraitPay\audit\2026-04-28-frontend-audit.docx'
md_to_docx(md_report, docx_report)

# Convert contract templates
base = r'I:\Portraitpay ai\Doc\合同模板设计\中文草稿'
for fname in sorted(os.listdir(base)):
    if fname.endswith('.md'):
        md_path = os.path.join(base, fname)
        docx_path = os.path.join(base, fname.replace('.md', '.docx'))
        md_to_docx(md_path, docx_path)

print("All done!")
